# -*- coding: utf-8 -*-
"""DOCX 表格单元格内图片输入契约：python-docx 回退时表格内图不再丢失。

根因回归：历史 _parse_docx 回退只遍历顶层段落图片，表格单元格内的图虽被上传却没有
锚点/markdown，导致 source_image_count=0。本测试构造"表格单元格内含图"的 docx，
强制走 python-docx 回退，断言该图 URL 进入最终 markdown。
"""

from io import BytesIO
from unittest import TestCase
from unittest.mock import patch

from docx import Document
from docx.shared import Inches
from PIL import Image

from services.document_parser_service import DocumentParserService, EmbeddedImageUploader


def _png_bytes(color=(200, 30, 30)) -> bytes:
    img = Image.new("RGB", (120, 120), color)
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _docx_with_table_image() -> bytes:
    doc = Document()
    doc.add_heading("标题", level=1)
    doc.add_paragraph("正文段落。")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "左单元格文字"
    # 在右单元格插入一张图片。
    run = table.cell(0, 1).paragraphs[0].add_run()
    run.add_picture(BytesIO(_png_bytes()), width=Inches(1.0))
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_with_merged_cell_image() -> bytes:
    """跨列合并单元格里含一张图：复现 python-docx row.cells 对 gridSpan
    重复返回同一 <w:tc>，导致同图被输出多次的 bug。"""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    # 第一行三个单元格合并成一个跨 3 列的大单元格，并在其中插入图片。
    merged = table.cell(0, 0).merge(table.cell(0, 1)).merge(table.cell(0, 2))
    run = merged.paragraphs[0].add_run()
    run.add_picture(BytesIO(_png_bytes()), width=Inches(1.0))
    table.cell(1, 0).text = "普通行"
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


class DocxTableImageTests(TestCase):
    def test_table_cell_image_enters_markdown(self):
        data = _docx_with_table_image()
        svc = DocumentParserService(enable_embedded_image_upload=True, enable_ocr=False, min_img_bytes=100)

        fake_url = "https://cdn/table_img_1.png"

        def fake_extract(cls_data, media_prefix, min_size=5120):
            # 返回一张“已抽取”的图，交给 upload 映射为公网 URL。
            return [("image1.png", _png_bytes(), "image/png")]

        def fake_upload(images):
            return {img[0]: fake_url for img in images}

        # 强制走 python-docx 回退（markitdown 返回空），并 mock 图片抽取/上传避免真实对象存储。
        with patch.object(svc, "_markitdown_convert", return_value=""), \
             patch.object(EmbeddedImageUploader, "extract_from_zip", classmethod(lambda cls, d, p, min_size=5120: [("image1.png", _png_bytes(), "image/png")])), \
             patch.object(EmbeddedImageUploader, "upload_images", classmethod(lambda cls, images: {img[0]: fake_url for img in images})):
            md = svc._parse_docx(data)

        self.assertIn(fake_url, md, f"表格单元格内图片 URL 未进入 markdown:\n{md}")

    def test_merged_cell_image_not_duplicated(self):
        """根因回归：跨列合并单元格内的图，在 markdown 里只出现一次，不因 gridSpan 重复。"""
        data = _docx_with_merged_cell_image()
        svc = DocumentParserService(enable_embedded_image_upload=True, enable_ocr=False, min_img_bytes=100)
        fake_url = "https://cdn/merged_img.png"

        with patch.object(svc, "_markitdown_convert", return_value=""), \
             patch.object(EmbeddedImageUploader, "extract_from_zip", classmethod(lambda cls, d, p, min_size=5120: [("image1.png", _png_bytes(), "image/png")])), \
             patch.object(EmbeddedImageUploader, "upload_images", classmethod(lambda cls, images: {img[0]: fake_url for img in images})):
            md = svc._parse_docx(data)

        self.assertEqual(
            md.count(fake_url), 1,
            f"合并单元格内图片被重复输出 {md.count(fake_url)} 次:\n{md}",
        )
